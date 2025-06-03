import os
import logging
import uuid
from docx import Document
from docx.shared import Pt, Inches
from app.schemas.paper import GeneratePaperRequest, JobStatus
from app.services.outline_service import OutlineService
from app.services.draft_service import DraftService
from app.services.document_service import DocumentService
from app.services.supabase_service import save_paper, upload_paper_file
# Configure logging
logger = logging.getLogger(__name__)

class PaperGeneratorService:
    """Service to handle paper generation logic"""
    
    def __init__(self):
        self.outline_service = OutlineService()
        self.draft_service = DraftService()
        self.output_dir = "outputs"
        os.makedirs(self.output_dir, exist_ok=True)
    
    async def generate_paper(self, job_id: str, request: GeneratePaperRequest, job_statuses: dict, user_id: str):
        """Generate a paper based on the provided details."""
        try:
            job_statuses[job_id] = {"status": "processing", "progress": 0, "message": "Starting paper generation"}
            
            # Generate paper content using our service pipeline
            paper_content = await self._generate_paper_content(request.paper_details)
            
            # Create document using DocumentService
            doc_service = DocumentService(
                title=paper_content["title"],
                content=paper_content["content"],
                sources=paper_content["sources"],
                student_name=request.student_details.student_name,
                professor_name=request.student_details.professor_name,
                class_name=request.student_details.class_name
            )
            
            file_path = os.path.join(self.output_dir, f"{job_id}.docx")
            doc_service.create_document(file_path)
            
            # Save paper metadata and file to Supabase
            paper_data = {
                "title": request.paper_details.title,
                "description": request.paper_details.description,
                "sections": request.paper_details.num_sections,
                "word_count": len(paper_content["content"].split()),
                "job_id": job_id
            }
            
            save_paper(user_id, paper_data, file_path)
            job_statuses[job_id] = {
                "status": "completed", 
                "file_path": file_path,
                "message": "Paper generation completed successfully"
            }
            
        except Exception as e:
            logger.error(f"Error generating paper: {str(e)}", exc_info=True)
            job_statuses[job_id] = {
                "status": "failed", 
                "error": str(e),
                "message": f"Failed to generate paper: {str(e)}"
            }
            
        return job_statuses[job_id]
    
    async def _generate_paper_content(self, paper_details):
        """Generate the content for the paper using our service pipeline."""
        try:
            # 1. Generate outline
            outline = self.outline_service.generate_outline(
                title=paper_details.title,
                description=paper_details.description,
                num_sections=paper_details.num_sections
            )
            
            # 2. Generate draft using outline and sources
            draft = self.draft_service.generate_draft(
                outline=outline,
                title=paper_details.title,
                vocab_level=paper_details.vocab_level
            )
            
            # The draft_service has already gathered sources during draft generation
            sources = self.draft_service.sources
            
            return {
                "title": paper_details.title,
                "content": draft,
                "sources": sources
            }
            
        except Exception as e:
            logger.error(f"Error in paper generation pipeline: {str(e)}", exc_info=True)
            raise
    
    def _format_as_docx(self, paper_content, paper_details, student_details, job_id):
        """Format the paper content as a Word document."""
        doc = Document()
        
        # Set document style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title page
        doc.add_paragraph(paper_content["title"]).alignment = 1  # Center alignment
        doc.add_paragraph(f"By: {student_details.student_name}").alignment = 1
        doc.add_paragraph(f"Course: {student_details.class_name}").alignment = 1
        doc.add_paragraph(f"Professor: {student_details.professor_name}").alignment = 1
        doc.add_paragraph(f"Date: {self._get_current_date()}").alignment = 1
        
        # Add page break after title page
        doc.add_page_break()
        
        # Add introduction
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph(paper_content["introduction"])
        
        # Add sections
        for section in paper_content["sections"]:
            doc.add_heading(section["title"], level=1)
            doc.add_paragraph(section["content"])
        
        # Add conclusion
        doc.add_heading("Conclusion", level=1)
        doc.add_paragraph(paper_content["conclusion"])
        
        # Save document
        file_path = os.path.join(self.output_dir, f"{job_id}.docx")
        doc.save(file_path)
        
        return file_path
    
    def _count_words(self, paper_content):
        """Count the total words in the paper."""
        word_count = 0
        
        # Count words in introduction
        word_count += len(paper_content["introduction"].split())
        
        # Count words in sections
        for section in paper_content["sections"]:
            word_count += len(section["content"].split())
        
        # Count words in conclusion
        word_count += len(paper_content["conclusion"].split())
        
        return word_count
    
    def _get_current_date(self):
        """Get the current date formatted for the paper."""
        from datetime import datetime
        return datetime.now().strftime("%B %d, %Y") 