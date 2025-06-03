from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import logging
import os
from app.models.source import Source

# Configure logging
logger = logging.getLogger(__name__)

class DocumentService:
    """Service for generating formatted documents"""
    
    def __init__(self, 
                 title: str,
                 content: str,
                 sources: list,
                 student_name: str,
                 professor_name: str,
                 class_name: str,
                 student_last_name: str = None):
        self.title = title
        self.content = content
        self.sources = sources
        self.student_name = student_name
        self.professor_name = professor_name
        self.class_name = class_name
        # If no last name provided, try to get it from the full name
        self.student_last_name = student_last_name or student_name.split()[-1]
        self.date = datetime.now().strftime("%d %B %Y")
    
    def create_page_number(self, paragraph):
        """Create a page number field at the end of the given paragraph."""
        try:
            run = paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
        except Exception as e:
            logger.error(f"Error creating page number: {str(e)}", exc_info=True)

    def create_document(self, output_path: str = "final_paper.docx"):
        """Create a formatted Word document"""
        try:
            logger.info(f"Creating document: {output_path}")
            
            # Create output directory if needed
            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
            
            # Create a completely fresh document
            doc = Document()
            
            # Set up document formatting
            for section in doc.sections:
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                
                # Remove ALL headers from first page
                section.different_first_page_header = True
                
                # Set up headers for subsequent pages only
                if section.header:
                    for paragraph in section.header.paragraphs:
                        paragraph.clear()  # Clear any existing content
                    
                    header_para = section.header.paragraphs[0]
                    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = header_para.add_run(f"{self.student_last_name} ")
                    self.create_page_number(header_para)
            
            # TITLE PAGE
            # Top right corner: Last name and page number
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run()
            run.font.size = Pt(12)
            
            # Add space before title
            for _ in range(4):
                p = doc.add_paragraph()
                p.space_after = Pt(0)
                p.space_before = Pt(0)
            
            # Title
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(self.title)
            run.font.size = Pt(12)
            
            # Add space before student info
            for _ in range(8):
                p = doc.add_paragraph()
                p.space_after = Pt(0)
                p.space_before = Pt(0)
            
            # Student info block
            info_elements = [
                self.student_name,
                self.class_name,
                self.professor_name,
                self.date
            ]
            
            for text in info_elements:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(text)
                run.font.size = Pt(12)
                para.space_after = Pt(12)
            
            # Create a new section for the content
            section_break = doc.add_paragraph()
            section_break.add_run().add_break(WD_BREAK.PAGE)
            
            # Process content
            paragraphs = [p for p in self.content.replace('\n\n', '\n').split('\n') if p.strip()]
            
            # Check if we need to add content directly to avoid blank page
            if len(paragraphs) > 0:
                # Add the first paragraph directly
                section_break.clear()
                section_break.paragraph_format.first_line_indent = Inches(0.5)
                section_break.paragraph_format.line_spacing = 2.0
                section_break.paragraph_format.space_after = Pt(0)
                section_break.paragraph_format.space_before = Pt(0)
                run = section_break.add_run(paragraphs[0].strip())
                run.font.size = Pt(12)
                # Remove the first paragraph since we've already added it
                paragraphs = paragraphs[1:]
            
            # Set up the style for remaining paragraphs
            style = doc.styles['Normal']
            style.paragraph_format.first_line_indent = Inches(0.5)
            style.paragraph_format.line_spacing = 2.0
            
            # Add the remaining paragraphs
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(style=style)
                    para.paragraph_format.first_line_indent = Inches(0.5)
                    para.paragraph_format.line_spacing = 2.0
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                    run = para.add_run(para_text.strip())
                    run.font.size = Pt(12)
            
            # Add References page if there are sources
            if self.sources:
                doc.add_page_break()
                
                # References header
                ref_header = doc.add_paragraph("References")
                ref_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ref_header.runs[0].font.size = Pt(12)
                ref_header.runs[0].font.bold = False
                ref_header.space_after = Pt(24)
                
                # Create hanging indent style for references
                ref_style = doc.styles.add_style('APA_Reference', WD_STYLE_TYPE.PARAGRAPH)
                ref_style.base_style = doc.styles['Normal']
                ref_style.font.size = Pt(12)
                ref_style.paragraph_format.first_line_indent = Inches(-0.5)  # Hanging indent
                ref_style.paragraph_format.left_indent = Inches(0.5)
                ref_style.paragraph_format.line_spacing = 2.0
                ref_style.paragraph_format.space_after = Pt(0)
                ref_style.paragraph_format.space_before = Pt(0)
                
                # Add each cited reference
                for source in self.sources:
                    ref_para = doc.add_paragraph(style='APA_Reference')
                    run = ref_para.add_run(source.apa_citation)
                    run.font.size = Pt(12)
                    ref_para.paragraph_format.first_line_indent = Inches(-0.5)
                    ref_para.paragraph_format.left_indent = Inches(0.5)
                    ref_para.paragraph_format.line_spacing = 2.0
            
            # Save the document
            doc.save(output_path)
            logger.info(f"Document created successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating document: {str(e)}", exc_info=True)
            raise 